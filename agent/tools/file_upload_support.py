# Source note: Adapted from agent/tools/file_upload_supprt.py in EBC

import base64
import logging
from io import BytesIO
from typing import Dict, List, Any, Optional
from docx import Document
from PyPDF2 import PdfReader
from google.adk.tools import ToolContext

logger = logging.getLogger(__name__)

# Constants
LOG_IDENTIFIER = "GTM_PP"
MAX_UPLOAD_SIZE_MB = 10
SUPPORTED_MIME_TYPES = [
    'application/pdf',
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
    'application/vnd.openxmlformats-officedocument.presentationml.presentation',
    'text/plain',
    'text/csv',
    'application/json',
]

def docx_bytes_to_all_text(data: bytes) -> str:
    """
    Extract all text from a DOCX file.
    Source note: Exact implementation from EBC agent/tools/file_upload_supprt.py
    """
    try:
        doc = Document(BytesIO(data))
        parts = []
        
        # Extract paragraphs
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text.strip())
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    parts.append(" | ".join(row_text))
        
        # Extract headers and footers
        for section in doc.sections:
            header = section.header
            footer = section.footer
            for p in header.paragraphs + footer.paragraphs:
                if p.text.strip():
                    parts.append(p.text.strip())
        
        return "\n".join(parts)
    except Exception as e:
        logger.error(f"[{LOG_IDENTIFIER}] Error extracting DOCX text: {str(e)}")
        return ""


def pdf_bytes_to_text(data: bytes) -> str:
    """
    Extract text from a PDF file.
    Additional handler not in original EBC but useful for GTM.
    """
    try:
        reader = PdfReader(BytesIO(data))
        text_parts = []
        for page_num, page in enumerate(reader.pages, 1):
            try:
                text = page.extract_text()
                if text:
                    text_parts.append(f"--- Page {page_num} ---\n{text}")
            except Exception as e:
                logger.warning(f"[{LOG_IDENTIFIER}] Failed to extract text from page {page_num}: {str(e)}")
        return "\n".join(text_parts)
    except Exception as e:
        logger.error(f"[{LOG_IDENTIFIER}] Error extracting PDF text: {str(e)}")
        return ""


async def list_artifacts(tool_context: ToolContext = None) -> Dict[str, Any]:
    """
    List and process all uploaded artifacts.
    ALWAYS WORKS - returns empty list if no uploads, no config needed.
    Source note: Adapted from EBC agent/tools/file_upload_supprt.py
    """
    if not tool_context:
        logger.error(f"[{LOG_IDENTIFIER}] No tool context provided to list_artifacts")
        return {"status": "error", "error": "Tool context required", "grounding_material": []}
    
    try:
        logger.info(f"[{LOG_IDENTIFIER}] Checking for uploaded artifacts...")
        names = await tool_context.list_artifacts()
        
        if not names:
            # No artifacts uploaded - this is perfectly normal
            logger.debug(f"[{LOG_IDENTIFIER}] No artifacts found - proceeding without uploads")
            return {
                "status": "success",
                "message": "No uploaded documents found",
                "grounding_material": []
            }
        
        # User has uploaded documents - process them
        logger.info(f"[{LOG_IDENTIFIER}] Found {len(names)} uploaded documents, processing...")
        artifacts_total = []
        errors = []
        
        for name in names:
            logger.info(f"[{LOG_IDENTIFIER}] Processing artifact: {name}")
            try:
                artifact_data = await get_artifact(name, tool_context=tool_context)
                if "error" not in artifact_data:
                    artifacts_total.append(artifact_data)
                    logger.info(f"[{LOG_IDENTIFIER}] Successfully processed: {name}")
                else:
                    errors.append({"filename": name, "error": artifact_data["error"]})
                    logger.warning(f"[{LOG_IDENTIFIER}] Error processing {name}: {artifact_data['error']}")
            except Exception as e:
                errors.append({"filename": name, "error": str(e)})
                logger.error(f"[{LOG_IDENTIFIER}] Exception processing {name}: {str(e)}")
        
        logger.info(f"[{LOG_IDENTIFIER}] Processed {len(artifacts_total)} artifacts successfully")
        
        result = {
            "status": "success",
            "message": f"Found and processed {len(artifacts_total)} uploaded documents",
            "grounding_material": artifacts_total
        }
        
        if errors:
            result["processing_errors"] = errors
        
        return result
        
    except Exception as e:
        logger.exception(f"[{LOG_IDENTIFIER}] Failed to list artifacts: {str(e)}")
        return {
            "status": "error",
            "error": str(e),
            "grounding_material": []
        }


async def get_artifact(filename: str, tool_context: ToolContext = None) -> Dict[str, Any]:
    """
    Get and process a single artifact.
    Source note: Enhanced version of EBC agent/tools/file_upload_supprt.py with better MIME handling
    """
    if not tool_context:
        return {"error": "Tool context required"}
    
    try:
        logger.debug(f"[{LOG_IDENTIFIER}] Loading artifact: {filename}")
        part = await tool_context.load_artifact(filename=filename)
        
        if not part:
            return {"error": "Artifact not found"}
        
        # Handle different artifact structures
        if isinstance(part, dict) and 'inlineData' in part:
            inline_data = part['inlineData']
            mime_type = inline_data.get('mimeType', 'unknown')
            
            # Validate MIME type
            if mime_type not in SUPPORTED_MIME_TYPES:
                logger.warning(f"[{LOG_IDENTIFIER}] Unsupported MIME type: {mime_type}")
                return {
                    "error": f"Unsupported file type: {mime_type}",
                    "filename": filename,
                    "mime_type": mime_type
                }
            
            # Extract data based on MIME type
            data = ''
            data_bytes = base64.b64decode(inline_data['data'])
            
            # Check file size
            size_mb = len(data_bytes) / (1024 * 1024)
            if size_mb > MAX_UPLOAD_SIZE_MB:
                return {
                    "error": f"File too large: {size_mb:.2f}MB (max: {MAX_UPLOAD_SIZE_MB}MB)",
                    "filename": filename,
                    "size_mb": size_mb
                }
            
            # Process based on MIME type
            if mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                data = docx_bytes_to_all_text(data_bytes)
            elif mime_type == 'application/pdf':
                data = pdf_bytes_to_text(data_bytes)
            elif mime_type in ['text/plain', 'text/csv', 'application/json']:
                data = data_bytes.decode('utf-8', errors='ignore')
            else:
                # For other supported types, try text extraction
                try:
                    data = data_bytes.decode('utf-8', errors='ignore')
                except:
                    data = f"Binary content ({mime_type})"
            
            if not data:
                return {
                    "error": "No content extracted",
                    "filename": filename,
                    "mime_type": mime_type
                }
            
            return {
                "filename": filename,
                "mime_type": mime_type,
                "size_bytes": len(data_bytes),
                "size_mb": size_mb,
                "data": data,
                "preview": data[:500] + "..." if len(data) > 500 else data
            }
        
        # Handle text-only artifacts
        elif hasattr(part, 'text') and part.text:
            return {
                "filename": filename,
                "mime_type": "text/plain",
                "size_bytes": len(part.text),
                "data": part.text,
                "preview": part.text[:500] + "..." if len(part.text) > 500 else part.text
            }
        
        else:
            return {
                "error": "Unknown artifact format",
                "filename": filename
            }
            
    except Exception as e:
        logger.error(f"[{LOG_IDENTIFIER}] Error loading artifact {filename}: {str(e)}")
        return {
            "error": str(e),
            "filename": filename
        }