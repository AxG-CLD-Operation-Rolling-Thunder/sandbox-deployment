from google.adk.tools import ToolContext
from io import BytesIO
from docx import Document
import base64

def docx_bytes_to_all_text(data: bytes) -> str:
    doc = Document(BytesIO(data))
    parts = []

    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())

    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                parts.append(" | ".join(row_text))

    for section in doc.sections:
        header = section.header
        footer = section.footer
        for p in header.paragraphs + footer.paragraphs:
            if p.text.strip():
                parts.append(p.text.strip())

    return "\n".join(parts)


async def list_artifacts(tool_context: ToolContext = None) -> dict:
    names = await tool_context.list_artifacts()
    artifacts_total = []

    for name in names:
        print(f"Found artifact: {name} ... adding to context")

        artifacts_total.append(await get_artifact(name, tool_context=tool_context))

    return {"grounding_material": artifacts_total}


async def get_artifact(filename: str, tool_context: ToolContext = None) -> dict:
    part = await tool_context.load_artifact(filename=filename)
    print(type(part))

    if part and isinstance(part, dict):
        print(part.keys())
        print(part.values())

    data = ""

    if not (part and part.get("inlineData")):
        return {"error": "not found"}

    if (
        part["inlineData"]["mimeType"]
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ):
        data = docx_bytes_to_all_text(base64.b64decode(part["inlineData"]["data"]))
    elif part["inlineData"]["mimeType"] == "text/plain":
        data = part["inlineData"]["data"]

    if not data:
        return {"error": "empty or invalid mime type"}

    return {
        "filename": filename,
        "mime_type": part["inlineData"]["mimeType"],
        "size_bytes": len(part["inlineData"]["data"]),
        "data": data,
    }
